from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "PID_dashboard_documentacion.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_doc(doc: Document) -> None:
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)

    for name, size, bold, color in [
        ("Title", 26, True, RGBColor(0, 0, 0)),
        ("Heading 1", 16, True, RGBColor(31, 78, 121)),
        ("Heading 2", 12.5, True, RGBColor(31, 78, 121)),
        ("Heading 3", 11, True, RGBColor(31, 78, 121)),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = color


def p(doc: Document, text: str, style: str | None = None, bold: bool = False, italic: bool = False, align=None):
    para = doc.add_paragraph(style=style)
    if align is not None:
        para.alignment = align
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Arial"
    return para


def bullet(doc: Document, text: str) -> None:
    p(doc, text, style="List Bullet")


def number(doc: Document, text: str) -> None:
    p(doc, text, style="List Number")


def add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Inches(2.1)
    table.columns[1].width = Inches(4.9)
    hdr = table.rows[0].cells
    hdr[0].text = "Campo"
    hdr[1].text = "Descripción"
    for c in hdr:
        set_cell_shading(c, "D9EAF7")
        for p_ in c.paragraphs:
            p_.runs[0].bold = True
    set_repeat_table_header(table.rows[0])
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
    doc.add_paragraph("")


def build() -> None:
    doc = Document()
    style_doc(doc)

    # Cover
    p(doc, "PID Dashboard", style="Title")
    p(doc, "Documentación técnica consolidada del modelo de datos, destinos y POIs", italic=True)
    p(doc, "Versión preparada para trabajo interno y explicación funcional del flujo completo.", italic=True)
    doc.add_paragraph("")
    p(doc, "Ámbito", bold=True)
    bullet(doc, "Modelo conceptual del proyecto PID.")
    bullet(doc, "Normalización de destinos, POIs y geografía.")
    bullet(doc, "Trazabilidad desde los ficheros originales hasta las salidas finales.")
    doc.add_page_break()

    p(doc, "1. Objetivo del documento", style="Heading 1")
    p(doc, "Este documento explica el flujo completo del proyecto: qué entra, cómo se normaliza, cómo se relacionan las entidades y qué salidas genera el script. La intención es que sirva como base documental estable para el visor, para análisis internos y para futuras automatizaciones.")

    p(doc, "2. Modelo conceptual", style="Heading 1")
    p(doc, "La lógica del PID se apoya en tres capas: destino, POI y geografía.")
    bullet(doc, "Destino: unidad de negocio turística. Puede ser municipio, comarca, diputación, consell o comunidad autónoma.")
    bullet(doc, "POI: punto de interés normalizado desde el JSON de seguimiento.")
    bullet(doc, "Geografía: jerarquía canónica de municipio, provincia y comunidad autónoma.")

    p(doc, "3. Entradas del sistema", style="Heading 1")
    p(doc, "Los ficheros fuente que alimentan la normalización son los siguientes:")
    bullet(doc, "JSON de seguimiento: docs/seguimiento_PRO_20260727.json")
    bullet(doc, "Excel de beneficiarios: docs/InfoBeneficiarios_v1.61_20260730.xlsx")
    bullet(doc, "Excel de gestores: docs/InfoBeneficiarios_Gestores_20260730.xlsx")
    bullet(doc, "Capas geográficas: municipios, provincias y comunidades autónomas en GeoJSON")

    p(doc, "4. Procesado de destinos", style="Heading 1")
    p(doc, "El pipeline convierte los Excel en tablas limpias, homogéneas y listas para cruce. La normalización conserva el nombre legible y añade claves normalizadas para evitar problemas de tildes, variantes lingüísticas o diferencias de escritura.")
    add_kv_table(doc, [
        ("uso_pid_normalizado.csv", "Resumen de uso de la PID por destino, con entidad gestora, códigos territoriales y número de usuarios."),
        ("modulos_comunes_normalizado.csv", "Módulos comunes y solución digital, con contexto territorial y fecha de contratación."),
        ("datos_turisticos_normalizado.csv", "Metadatos de grafo y seguimiento de tripletas por destino."),
        ("datos_no_ontologicos_normalizado.csv", "Datos textuales o no ontológicos resumidos para análisis."),
        ("destinos_gestores_normalizado.csv", "Relación entre destino gestor y destino gestionado."),
    ])

    p(doc, "5. Procesado de POIs", style="Heading 1")
    p(doc, "El JSON de seguimiento contiene 22.781 puntos. El script normaliza los POIs a una tabla de trabajo, resuelve el destino asociado mediante dti, y después añade el enriquecimiento geográfico.")
    bullet(doc, "POI normalizado: identificación, clase, nombre, destino, coordenadas y trazabilidad.")
    bullet(doc, "POI enriquecido: asignación territorial por municipio, provincia y CCAA.")
    bullet(doc, "Estados de match geográfico: matched, nearest y unmatched.")

    p(doc, "6. Diccionario funcional de campos", style="Heading 1")
    p(doc, "Las columnas finales de `pois_enriquecidos.csv` se agrupan así:")
    add_kv_table(doc, [
        ("Identidad", "phase, entity_key, entity_name, entity_type, category, name, class, score y razon."),
        ("Coordenadas", "lat, lon, has_coordinates y coord_source."),
        ("Destino", "dti, dti_norm, uri, range2_out y los campos destino_*."),
        ("Geografía", "codigo_ine, nombre_municipio, codigo_provincia, nombre_provincia, codigo_ccaa y nombre_ccaa."),
        ("Estado", "match_status_geo para distinguir matched, nearest y unmatched."),
    ])

    p(doc, "7. Relación entre destino y geografía", style="Heading 1")
    p(doc, "El destino no siempre coincide con la jerarquía territorial pura. Hay entidades que son agregados turísticos y no tienen código INE propio, como comarcas. Por eso el modelo mantiene separada la capa de destino y la capa geográfica, pero las conecta siempre que sea posible.")
    bullet(doc, "Municipio: unidad mínima territorial.")
    bullet(doc, "Provincia y CCAA: capas superiores para agregación y análisis.")
    bullet(doc, "Comarca, diputación y consell: entidades turísticas que pueden no encajar de forma directa con el código INE.")

    p(doc, "8. Capa comarcal y propuesta municipio-comarca", style="Heading 1")
    p(doc, "Se ha incorporado una capa comarcal propia y una propuesta inferida de pertenencia municipio-comarca. La propuesta se construye a partir de menciones textuales del audit y se considera válida como base de trabajo hasta revisión experta.")
    bullet(doc, "comarcas_lite.csv: inventario de comarcas o entidades comarcales detectadas.")
    bullet(doc, "municipio_comarca_propuesta.csv: relación inferida municipio-comarca con peso, orden y confianza.")
    bullet(doc, "comarcas_propuesta.csv: resumen por comarca de la propuesta de relación.")

    p(doc, "8.1 Pendientes de validación experta", style="Heading 2")
    bullet(doc, "Comarcas con nombres muy parecidos a municipios y posibles falsos positivos.")
    bullet(doc, "Menciones textuales tomadas de descripciones largas que requieran comprobación humana.")
    bullet(doc, "Entidades supramunicipales que el audit trate como comarca pero que funcionen como consell, mancomunidad o federación.")
    bullet(doc, "Registros sin provincia explícita o con provincia ambigua.")
    bullet(doc, "Estado operativo inicial: propuesto. Si se valida, pasará a validado; si no, se corregirá o eliminará.")

    p(doc, "9. Flujo extremo a extremo", style="Heading 1")
    number(doc, "Entran los Excel de beneficiarios y gestores, el JSON de seguimiento y las capas geográficas.")
    number(doc, "El script detecta estructura, no nombre exacto del archivo, y normaliza todo a CSV reproducibles.")
    number(doc, "Se resuelven destinos por nombre normalizado y POIs por dti normalizado.")
    number(doc, "Se asigna municipio, provincia y comunidad autónoma mediante cruce espacial.")
    number(doc, "Se generan las salidas finales en normalized/destinos, normalized/pois y normalized/geo.")

    p(doc, "10. Decisiones de diseño", style="Heading 1")
    bullet(doc, "No se mezcla la lógica del destino con la lógica geográfica.")
    bullet(doc, "Se conserva trazabilidad suficiente para auditar el cruce.")
    bullet(doc, "Los registros sin match se dejan marcados, no se fuerzan.")
    bullet(doc, "La normalización está pensada para ejecutarse con un solo comando.")

    p(doc, "11. Fuentes y documentación relacionada", style="Heading 1")
    bullet(doc, "docs/pipeline-normalizacion.md")
    bullet(doc, "docs/destinos-normalizados-columnas.md")
    bullet(doc, "docs/pois-enriquecidos-columnas.md")
    bullet(doc, "docs/modelo-relacional-pid.md")
    bullet(doc, "docs/inventario-final-salidas.md")
    bullet(doc, "scripts/normalize_pid_data.py")

    doc.save(OUT)


if __name__ == "__main__":
    build()
